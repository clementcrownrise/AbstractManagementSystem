from django.shortcuts import render, redirect, get_object_or_404
from .forms import ArticleForm, CasereportForm
from django.contrib  import messages
from .models import Article, ArticleReviewer, Comment, Casereport, ReviewerReport
from accounts.models import Account
from conference.models import Conference
from reviewers.models import Reviewer
from faculty.models import Faculty
from django.core.mail import EmailMessage
from django.template.loader import render_to_string
from authors.models import Authors
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from docx import Document


@login_required
def create_casereport(request):
    if request.method == 'POST':
        casereportform = CasereportForm(request.POST)
        print('I am getting to casereport')
        if casereportform.is_valid():
            casereport = casereportform.save(commit=False)
            casereport.user = request.user
            casereport.status = 'Pending'
            casereport.save()
            

            fullnames = request.POST.getlist("fullname[]")
            emails = request.POST.getlist("email[]")
            phones = request.POST.getlist("phone[]")
            affiliations = request.POST.getlist("affiliation[]")

            for fullname, email, phone, affiliation in zip(
                fullnames, 
                emails, 
                phones, 
                affiliations
            ):
                Authors.objects.create(
                    casereport=casereport,   # 
                    fullname=fullname,
                    email=email,
                    phone=phone,
                    affiliation=affiliation
                    )
            messages.success(request, 'Abstract Submitted Successfully')
            #I need to send a mail here
            mail_subject = 'Abstract Submission Notification'
            to_email = set()
            to_email.add('clementcrownrise@gmail.com')
            to_email.add('wacs.scicom@gmail.com')
            to_email.add('eaameh@gmail.com')
            to_email.add(casereport.user.email) 
            message = render_to_string('article/abstractSubmissionEmail.html',{
                'casereport':casereport,
            })
            send_email = EmailMessage(mail_subject, message, to=list(to_email))
            send_email.content_subtype='html'
            send_email.send()
           #i will send to admin below 
            to_emailadmin = set()
            to_emailadmin.add('clementcrownrise@gmail.com')
            to_emailadmin.add('eaameh@gmail.com')
            to_emailadmin.add('wacs.scicom@gmail.com')
            messageadmin = render_to_string('article/abstractSubmissionEmailFacultyHead.html',{
                'casereport':casereport,
            })
            send_emailadmin = EmailMessage(mail_subject, messageadmin, to=list(to_emailadmin))
            send_emailadmin.content_subtype='html'
            send_emailadmin.send()

            return redirect("dashboard")
    else:
        casereportform = CasereportForm()
        print('something is wrong here')
        print(casereportform.errors)



    context ={'casereportform':casereportform,}
    return render(request, 'article/create_casereport.html', context)


@login_required
def facultylisting(request, pk):
    faculty = get_object_or_404(Faculty, pk=pk)
    articles = Article.objects.filter(faculty = faculty).prefetch_related('reviewer_reports')
    casereports = Casereport.objects.filter(faculty=faculty).prefetch_related('reviewer_reports')

    context ={'faculty':faculty,
              'articles':articles,
              'casereports':casereports}
    return render(request, 'accounts/admin/faclisting.html',context )


# Create your views here.
@login_required
def create_abstract(request):

    if request.method == 'POST':
        originalform = ArticleForm(request.POST)
        print('IT is reaching')
        if originalform.is_valid():
                article = originalform.save(commit=False)
                article.user = request.user
                article.status = 'Pending'
                article.save()
                print(request.user)
            

                fullnames = request.POST.getlist("fullname[]")
                emails = request.POST.getlist("email[]")
                phones = request.POST.getlist("phone[]")
                affiliations = request.POST.getlist("affiliation[]")

                for fullname, email, phone, affiliation in zip(
                        fullnames, 
                        emails, 
                        phones, 
                        affiliations
                    ):
                    Authors.objects.create(
                        article=article,   # 
                        fullname=fullname,
                        email=email,
                        phone=phone,
                        affiliation=affiliation
                    )
                messages.success(request, 'Abstract Submitted Successfully')
                mail_subject = 'Abstract Submission Notification'
                to_email = set()
                to_email.add('clementcrownrise@gmail.com')
                to_email.add('wacs.scicom@gmail.com')
                to_email.add('eaameh@gmail.com')
                to_email.add(article.user.email)
                message = render_to_string('article/abstractSubmissionEmail.html',{
                'article':article,
                     })
                send_email = EmailMessage(mail_subject, message, to=list(to_email))
                send_email.content_subtype='html'
                send_email.send()

                #i also waant to send another email blow to chairman
                #i will send to admin below 
                to_emailadmin = set()
                to_emailadmin.add('clementcrownrise@gmail.com')
                to_emailadmin.add('eaameh@gmail.com')
                to_emailadmin.add('wacs.scicom@gmail.com')
                messageadmin = render_to_string('article/abstractSubmissionEmailFacultyHead.html',{
                    'article':article,
                })
                send_emailadmin = EmailMessage(mail_subject, messageadmin, to=list(to_emailadmin))
                send_emailadmin.content_subtype='html'
                send_emailadmin.send()

                return redirect("dashboard")
        

            
        else:
            print(originalform.errors)

    else:
        originalform = ArticleForm()
        print('something is wrong here')
        print(originalform.errors)
                  
    context ={'originalform':originalform,
              }
    return render(request, 'article/create_abstract.html', context)








@login_required
def view_abstract(request, id):
    #print(id)
    revieweremails = []
    user = request.user

    #I want to fetch the result of this article if it has any 
    existingResults = ReviewerReport.objects.filter(article_id = id)
          


    if user.user_type == 'candidate':

        article = get_object_or_404(Article, id=id, user=request.user)
        print(article.faculty)

    
    elif user.user_type == 'admin' and not user.faculty:
        #I only get the reviewers of the viewed abstract user faculty
        article = get_object_or_404(Article, id=id)
        revieweremails = Reviewer.objects.filter(
            faculty = article.faculty).values_list('id', 'email')
        
    elif user.user_type == 'admin' and user.faculty:
        #I only get the reviewers of the logged in Admin user faculty
        revieweremails = Account.objects.filter(
            faculty = request.user.faculty, user_type = 'reviewer').values_list('id', 'email')
        article = get_object_or_404(Article, id=id, user__faculty = request.user.faculty)
    else:
        messages.error(request, 'You can NOT view an article you do not own')
        return redirect('dashboard')
    comments = article.comments.all().order_by('-created_at')
    #i need to get all currently assigned reviewers
    assigned_reviewers = ArticleReviewer.objects.filter(article=article).select_related('reviewer')
    context={
        'article':article,
        'revieweremails':revieweremails,
        'comments':comments,
        'existingResults':existingResults,
        'assigned_reviewers':assigned_reviewers,
    }
    
    return render(request, 'article/view_abstract.html', context)


@login_required
def view_casereport(request, id):
    #print(id)
    revieweremails = []
    
    user = request.user
     #I want to fetch the result of this article if it has any 
    existingResults = ReviewerReport.objects.filter(casereport_id = id)
    #print(vars(existingResults))
         
    if user.user_type == 'candidate':

        casereport = get_object_or_404(Casereport, id=id, user=request.user)
        

    elif user.user_type == 'reviewer':
        #article = get_object_or_404(Article, id=id, user__faculty = request.user.faculty)
        assigment = get_object_or_404(ArticleReviewer, article_id = id, reviewer = request.user)
        casereport = assigment.article
    elif user.user_type == 'admin' and not user.faculty:
        #I only get the reviewers that belongs to the viewed case study.
        casereport = get_object_or_404(Casereport, id=id)
        revieweremails = Reviewer.objects.filter(
            faculty = casereport.faculty).values_list('id', 'email')
    elif user.user_type == 'admin' and user.faculty:
        #I only get the reviewers of the logged in Admin user faculty
        revieweremails = Account.objects.filter(
            faculty = request.user.faculty, user_type = 'reviewer').values_list('id', 'email')
        casereport = get_object_or_404(Casereport, id=id, user__faculty = request.user.faculty)
    else:
        messages.error(request, 'You can NOT view an article you do not own')
        return redirect('dashboard')
    comments = casereport.comments.all().order_by('-created_at')
    #i need to get all currently assigned reviewers
    assigned_reviewers = ArticleReviewer.objects.filter(casereport=casereport).select_related('reviewer')
    context={
        'casereport':casereport,
        'revieweremails':revieweremails,
        'comments':comments,
        'existingResults':existingResults,
        'assigned_reviewers':assigned_reviewers,
    }
    
    return render(request, 'article/view_casereport.html', context)


@login_required
def assign_reviewer(request,pk):
    #return HttpResponse( 'I am getting here')
    article  = Article.objects.get(id=pk)
    if request.method == 'POST':
        reviewer_id = request.POST.get('reviewer')
        #print(reviewer_id)
        if not reviewer_id:
            messages.error(request, "Failure: No reviewer was selected.")
            return redirect(request.META.get('HTTP_REFERER', '/'))
        #check if the revieer is already assigned 
        already_assigned = ArticleReviewer.objects.filter(
            article=article, 
            reviewer_id=reviewer_id
        ).exists()
        
        if not already_assigned:

            assignment = ArticleReviewer.objects.create(
                article = article,
                reviewer_id = reviewer_id
            )
            review_link =request.build_absolute_uri(
                f"/abstracts/review/{assignment.token}/"
            )
            #I need to send a message to send a mail to the reviewer
            mail_subject = 'Abstract Assignment Notification'
            message = render_to_string('article/ReviewerAssignmentEmail.html',{
                'article':article,
                'review_link':review_link
            })
            to_email = assignment.reviewer.email
            print(to_email)
            send_email = EmailMessage(mail_subject, message, to=[to_email])
            send_email.content_subtype= 'html'
            send_email.send()
        
            messages.success(request,'Reviewer has been assigned successfully')
        #print(reviewer_id,article)
        else:
            messages.error(request, 'Error adding a reviewer, This reviewer is already added to this Abstract')

    return redirect(request.META.get('HTTP_REFERER','/'))


@login_required
def casereportassign_reviewer(request,pk):
    #return HttpResponse( 'I am getting here')
    casereport  = Casereport.objects.get(id=pk)
    #print(casereport)
    if request.method == 'POST':
        reviewer_id = request.POST.get('reviewer')
        #print("reviewer id :")
        if not reviewer_id:
            messages.error(request, "Failure: No reviewer was selected.")
            return redirect(request.META.get('HTTP_REFERER', '/'))
        #check if the revieer is already assigned 
        already_assigned = ArticleReviewer.objects.filter(
            casereport=casereport, 
            reviewer_id=reviewer_id
        ).exists()
        
        if not already_assigned:

            assignment = ArticleReviewer.objects.create(
                casereport = casereport,
                reviewer_id = reviewer_id
            )
            review_link = request.build_absolute_uri(
                f"/abstracts/reviewcasereport/{assignment.token}/"
            )
            #I need to send a message to send a mail to the reviewer
            mail_subject = 'Abstract Assignment Notification'
            message = render_to_string('article/ReviewerAssignmentEmail.html',{
                'casereport':casereport,
                'review_link':review_link
            })
            to_email = assignment.reviewer.email
            send_email = EmailMessage(mail_subject, message, to=[to_email])
            send_email.content_subtype = 'html'
            send_email.send()
        
            messages.success(request,'Reviewer has been assigned successfully')
        #print(reviewer_id,article)
        else:
            messages.error(request, 'Error adding a reviewer, ' \
            'This reviewer is already added to this Abstract')

    return redirect(request.META.get('HTTP_REFERER','/'))


#reviewArticle
def review_article(request, token):
    assigment = get_object_or_404(ArticleReviewer, token = token)
    #print(assigment)
    article = assigment.article
    #check existing record
    existingResults = ReviewerReport.objects.filter(reviewer_id=assigment.reviewer_id,
                                                   article_id = article.id)
    return render(request, 'reviews/review.html',{'article':article,
                                                  'existingResults':existingResults,
                                                      'assignment':assigment})

def review_casereport(request, token):
    assigment = get_object_or_404(ArticleReviewer, token = token)
    casereport = assigment.casereport
    existingResults = ReviewerReport.objects.filter(reviewer_id=assigment.reviewer_id,
                                                   casereport_id = casereport.id)
    return render(request, 'reviews/casereportreview.html',{'casereport':casereport,
                                                            'existingResults':existingResults,
                                                      'assignment':assigment})



#adminComment
@login_required
def admin_comment(request,pk):
    article = Article.objects.get(id=pk)
    if request.method =='POST':
        #print('I am getting here as casereport')
        status = request.POST.get('status')
        article.status= request.POST.get('status')
        article.save()
        comment = request.POST.get('comment')
        saved_comment = Comment.objects.create(
            article = article,
            status = status,
            comment = comment,
            user = request.user
        )
        if saved_comment:
            #I need to send a mail to all the users in this article

            mail_subject = 'Abstract Update Notification'
            message = render_to_string('article/commentNotification.html',{
                'article':article,
            })
            emails =set()
            #candidate's email
            if request.user.email != article.user.email:
                emails.add(article.user.email)
            #faculty chairman email
            if request.user.user_type != 'admin':
                emails.update(Account.objects.filter(
                    faculty=article.user.faculty,
                    user_type ='admin'
                ).values_list('email',flat=True))
            #reviewer's email
            emails.update(
                ArticleReviewer.objects.filter(article=article).values_list(
                    'reviewer__email', flat=True)
            )
            send_email = EmailMessage(mail_subject, message, to=list(emails))
            #print(emails)
            send_email.content_subtype = 'html'
            send_email.send()

            messages.success(request, 'Your comment was saved successfully')
        else:
            messages.error(request, 'Messages could not be saved')

    return redirect(request.META.get('HTTP_REFERER', '/'))

def reviewerResult(request, pk):
    #print('I am about to add result here')
    
    article = Article.objects.get(id=pk)
    casereport = Casereport.objects.get(id=pk)
    if request.method == 'POST':

        formtype = request.POST.get('formtype')
        if formtype == 'casereport':
            token = request.POST.get('token')
            reviewer = request.POST.get('reviewer_id')
            #print(reviewer)
            titlescore = request.POST.get('titlescore')
            introductionscore = request.POST.get('introductionscore')
            resultscore = request.POST.get('resultscore')
            conclusionscore = request.POST.get('conclusionscore')
            overallscore = request.POST.get('overallscore')
            totalscore = request.POST.get('totalscore')
            comments = request.POST.get('comment')
            #i need to check for existing record before saving
            existingResult = ReviewerReport.objects.filter(
                reviewer_id=reviewer,casereport_id = casereport.id) 
            if existingResult:
                messages.error(request, 
                            'Messages could not be saved you already provided report for this abstract')
            else:
                saved_result = ReviewerReport.objects.create(
                    reviewer_id = reviewer,
                    casereport_id = casereport.id,
                    titlescore = titlescore,
                    introductionscore = introductionscore,
                    methodscore = 0,
                    resultscore = resultscore,
                    conclusionscore = conclusionscore,
                    totalscore = totalscore,
                    overallscore = overallscore,
                    comments = comments
                )
            if saved_result:
                #I need to send a mail to all the users in this article

                mail_subject = 'Abstract Update Notification'
                message = render_to_string('article/commentNotification.html',{
                    'casereport':casereport,
                    'article':article
                })
                emails =set()
                emails.add(casereport.user.email)
                emails.add('clementcrownrise@gmail.com')
                emails.add('eaameh@gmail.com')
                emails.add('wacs.scicom@gmail.com')
                #candidate's email
                send_email = EmailMessage(mail_subject, message, to=list(emails))
                #print(emails)
                send_email.send()

                messages.success(request, 'Your comment was saved successfully')
            else:
                messages.error(request, 'Messages could not be saved')           

        elif formtype == 'article':

            token = request.POST.get('token')
            reviewer = request.POST.get('reviewer_id')
            #print(reviewer)
            titlescore = request.POST.get('titlescore')
            introductionscore = request.POST.get('introductionscore')
            methodscore = request.POST.get('methodscore')
            resultscore = request.POST.get('resultscore')
            conclusionscore = request.POST.get('conclusionscore')
            overallscore = request.POST.get('overallscore')
            totalscore = request.POST.get('totalscore')
            comments = request.POST.get('comment')
            #i need to check for existing record before saving
            existingResult = ReviewerReport.objects.filter(
                reviewer_id=reviewer,article_id = article.id)
            if existingResult:
                messages.error(request, 
                            'Messages could not be saved you already provided report for this abstract')
            else:
                saved_result = ReviewerReport.objects.create(
                    reviewer_id = reviewer,
                    article_id = article.id,
                    titlescore = titlescore,
                    introductionscore = introductionscore,
                    methodscore = methodscore,
                    resultscore = resultscore,
                    conclusionscore = conclusionscore,
                    totalscore = totalscore,
                    overallscore = overallscore,
                    comments = comments
                )
            if saved_result:
                #I need to send a mail to all the users in this article

                mail_subject = 'Abstract Update Notification'
                message = render_to_string('article/commentNotification.html',{
                    'casereport':casereport,
                    'article':article
                })
                emails =set()
                emails.add(article.user.email)
                emails.add('clementcrownrise@gmail.com')
                emails.add('eaameh@gmail.com')               
                emails.add('wacs.scicom@gmail.com')
                #candidate's email
                send_email = EmailMessage(mail_subject, message, to=list(emails))
                #print(emails)
                send_email.send()

                messages.success(request, 'Your comment was saved successfully')
            else:
                messages.error(request, 'Messages could not be saved')

    return redirect(request.META.get('HTTP_REFERER', '/'))

#adminComment
@login_required
def admin_casereportcomment(request,pk):
    #print('i am getting here')
    casereport = Casereport.objects.get(id=pk)
    if request.method =='POST':
        #print('I am getting here')
        status = request.POST.get('status')
        casereport.status = request.POST.get('status')
        casereport.save()
        comment = request.POST.get('comment')
        saved_comment = Comment.objects.create(
            casereport = casereport,
            status = status,
            comment = comment,
            user = request.user
        )
        if saved_comment:
            #I need to send a mail to all the users in this article

            mail_subject = 'Abstract Update Notification'
            message = render_to_string('article/commentNotification.html',{
                'casereport':casereport,
            })
            emails =set()
            #candidate's email
            emails.add('clement')
            if request.user.email != casereport.user.email:
                emails.add(casereport.user.email)
            #faculty chairman email
            if request.user.user_type != 'admin':
                emails.update(Account.objects.filter(
                    faculty=casereport.user.faculty,
                    user_type ='admin'
                ).values_list('email',flat=True))
            #reviewer's email
            emails.update(
                ArticleReviewer.objects.filter(casereport=casereport).values_list(
                    'reviewer__email', flat=True)
            )
            send_email = EmailMessage(mail_subject, message, to=list(emails))
            #print(emails)
            send_email.send()

            messages.success(request, 'Your comment was saved successfully')
        else:
            messages.error(request, 'Messages could not be saved')

    return redirect(request.META.get('HTTP_REFERER', '/'))



@login_required
def remove_reviewer(request, article_id, reviewer_id):
    ArticleReviewer.objects.filter(
        article_id=article_id,
        reviewer_id = reviewer_id
    ).delete()  
    messages.error(request, 'The reviewer has been removed succefully')
    return redirect(request.META.get('HTTP_REFERER', '/'))

@login_required
def remove_reviewercasereport(request, casereport_id, reviewer_id):
    ArticleReviewer.objects.filter(
        casereport_id=casereport_id,
        reviewer_id = reviewer_id
    ).delete()  
    messages.error(request, 'The reviewer has been removed succefully')
    return redirect(request.META.get('HTTP_REFERER', '/'))


@login_required
def exportall(request):
    conferences = Conference.objects.all()
    faculties = Faculty.objects.all()

    if request.method == 'POST':
        conference  = request.POST.get('conference')
        faculty = request.POST.get('faculty')
        status = request.POST.get('status')
        typ = request.POST.get('typ')
        document = Document()
        if typ == 'casereport':
            casereports = Casereport.objects.filter(abstracttype='casereport',
                                                    conference_id=conference, 
                                                    faculty_id = faculty,
                                                    status = status
                                                    )
            for casereport in casereports:
                document.add_heading(casereport.title, level=1)
                document.add_paragraph(f'Author:{casereport.user}')
                document.add_paragraph(f'Presenter Name:{casereport.presentername}')
                document.add_paragraph(f'Presenter Email:{casereport.presenteremail}')
                document.add_paragraph(f'Presenter Phone:{casereport.presenterphone}')
                for author in casereport.authors.all():
                    document.add_paragraph(
                        f"{author.fullname }, {author.email }, {author.phone }, ({author.affiliation})"
                    )
                document.add_paragraph(f'Case Report:{casereport.casereport}')
                document.add_paragraph(f'Keywords:{casereport.keywords}')

                document.add_paragraph(f'French Title :{casereport.frtitle}')
                document.add_paragraph(f'French Casereport:{casereport.frcasereport}')
                document.add_paragraph(f'French Keywords:{casereport.frkeywords}')
                document.add_paragraph(f'French Title :{casereport.frtitle}')

                document.add_paragraph(f'Status:{casereport.status}')
                document.add_paragraph(f'Conference:{casereport.conference}')

                document.add_paragraph(f'Faculty :{casereport.faculty}')
                document.add_page_break
            response = HttpResponse(
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            response['Content-Disposition'] = 'attachment; filename=articles.docx'
            document.save(response)
            return response
        
        if typ == 'originalreview':
            articles = Article.objects.filter(abstracttype='original',
                                                    conference_id=conference, 
                                                    faculty_id = faculty,
                                                    status = status
                                                    )
            for article in articles:
                document.add_heading(article.title, level=1)
                document.add_paragraph(f'Author:{article.user}')
                document.add_paragraph(f'Presenter Name: {article.presentername}')
                document.add_paragraph(f'Presenter Email: {article.presenteremail}')
                document.add_paragraph(f'Presenter Phone: {article.presenterphone}')
                for author in article.authors.all():
                    document.add_paragraph(
                        f"{author.fullname }, {author.email }, {author.phone }, ({author.affiliation})"
                    )
                document.add_paragraph(f'Introduction:{article.introduction}')
                document.add_paragraph(f'Methods: {article.methods}')
                document.add_paragraph(f'Results: {article.results}')
                document.add_paragraph(f'Limitations: {article.limitations}')
                document.add_paragraph(f'Conclusion: {article.conclusion}')

                document.add_paragraph(f'Keywords: {article.keywords}')

                document.add_paragraph(f'French Title : {article.frtitle}')
                document.add_paragraph(f'French Introduction: {article.frintroduction}')
                document.add_paragraph(f'French Methods: {article.frmethods}')
                document.add_paragraph(f'French Results: {article.frresults}')
                document.add_paragraph(f'French Results: {article.frlimitations}')
                document.add_paragraph(f'French Conclusion: {article.frconclusion}')

                document.add_paragraph(f'French Keywords: {article.frkeywords}')

                document.add_paragraph(f'Status: {article.status}')
                document.add_paragraph(f'Conference: {article.conference}')

                document.add_paragraph(f'Faculty : {article.faculty}')
                document.add_page_break
            response = HttpResponse(
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            response['Content-Disposition'] = 'attachment; filename=articles.docx'
            document.save(response)
            return response


    context = {'conferences':conferences,
               'faculties':faculties
               }
    return render(request, 'article/export.html',context)


